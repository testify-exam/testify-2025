import os
import bcrypt
import random
import smtplib
import textwrap 
from fastapi.staticfiles import StaticFiles
from fpdf import FPDF
from docx import Document
from zoneinfo import ZoneInfo
from datetime import datetime
from dotenv import load_dotenv
from pydantic import BaseModel
from docx.shared import Inches
from functools import lru_cache
from fastapi.templating import Jinja2Templates
from fastapi import FastAPI, Request, Form, Depends, HTTPException
from fastapi import FastAPI, Query, Request, Form, HTTPException,status
from fastapi.responses import HTMLResponse, RedirectResponse, JSONResponse, FileResponse
from typing import Any, Dict, List, Optional
from tempfile import NamedTemporaryFile
from starlette.middleware.sessions import SessionMiddleware
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from sqlalchemy import delete, insert, or_, select,func, update, and_
from sqlalchemy.exc import NoResultFound
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload, joinedload
from Database import *



load_dotenv()

EMAIL_USER = os.getenv("EMAIL_USER")
EMAIL_PASS = os.getenv("EMAIL_PASS")

import StudentModule,TeacherModule

app = FastAPI()

app.mount("/static", StaticFiles(directory="static"), name="static")

StudentModule.init(app)
TeacherModule.init(app)
@app.api_route("/health", methods=["GET", "HEAD"])
async def health_check():
    return {"status": "ok"}

app.add_middleware(SessionMiddleware, secret_key="jk23kjl2j3kj23l2giy32tf43ft432ti4f3223y4yg32u43vbvj")
templates = Jinja2Templates(directory="templates")

def require_login(request: Request):
    if not request.session.get("user_id"):
        raise HTTPException(status_code=401, detail="Please log in to access this page.")




@app.get("/", response_class=HTMLResponse)
async def root(request: Request):
    if request.session.get("user_id"):
        role = request.session.get("role")
        if role == "teacher":
            return RedirectResponse(url="/teacher/home", status_code=302)
        elif role == "student":
            return RedirectResponse(url="/student/home", status_code=302)
        elif role == "admin":
            return RedirectResponse(url="/admin/home", status_code=302)
    # Otherwise, show the login page
    return templates.TemplateResponse("login.html", {"request": request})

@app.get("/favicon.ico", include_in_schema=False)
async def favicon():
    return FileResponse("static/favicon.ico")


@app.get("/login", response_class=HTMLResponse)
async def login_get(request: Request):
    user_id = request.session.get("user_id")
    if user_id:
        role_redirects = {
            "teacher": "/teacher/home",
            "student": "/student/home",
            "admin": "/admin/home",
        }
        return RedirectResponse(url=role_redirects.get(request.session.get("role"), "/login"), status_code=302)

    return templates.TemplateResponse("login.html", {"request": request})

@app.post("/login", response_class=HTMLResponse)
async def login_post(
    request: Request,
    email: str = Form(...),
    password: str = Form(...),
    remember: bool = Form(False),
    db: AsyncSession = Depends(get_db) 
):
    try:
        result = await db.execute(
            select(User).where(User.email == email)
        )
        user = result.scalar_one_or_none()

        if not user:
            return templates.TemplateResponse("login.html", {
                "request": request,
                "error": "Invalid email or password."
            })

        # Verify password
        stored_hash = user.password_hash.encode('utf-8')
        if not bcrypt.checkpw(password.encode('utf-8'), stored_hash):
            return templates.TemplateResponse("login.html", {
                "request": request,
                "error": "Invalid email or password."
            })

        if user.approved == -1:
            return templates.TemplateResponse("login.html", {
                "request": request,
                "error": "Your account is blocked by admin."
            })
        if user.approved == 0:
            return templates.TemplateResponse("login.html", {
                "request": request,
                "error": "Your account is pending admin approval."
            })

        if remember:
            request.session["user_id"] = user.id
            request.session["role"] = user.role
            request.session["username"] = user.username  
        else:
            request.session.clear()

        # Role-based redirection
        redirect_map = {
            "teacher": "/teacher/home",
            "student": "/student/home",
            "admin": "/admin/home"
        }
        
        redirect_url = redirect_map.get(user.role)
        if not redirect_url:
            raise HTTPException(status_code=400, detail="Invalid user role.")

        return RedirectResponse(url=redirect_url, status_code=status.HTTP_302_FOUND)

    except Exception as e:
        # In production, log this properly (don't expose to user)
        print(f"Login error: {e}")  # Replace with proper logging
        return templates.TemplateResponse("login.html", {
            "request": request,
            "error": "An unexpected error occurred. Please try again."
        })


@app.get("/logout")
async def logout(request: Request):
    request.session.clear()
    return RedirectResponse(url="/login", status_code=302)


@lru_cache(maxsize=1)
async def get_departments_cached(db: AsyncSession):
    result = await db.execute(select(Department.id, Department.name).order_by(Department.name))
    return [{"id": dept.id, "name": dept.name} for dept in result.all()]

@app.get("/register", response_class=HTMLResponse)
async def register_get(
    request: Request,
    db: AsyncSession = Depends(get_db)
):
    departments = await get_departments_cached(db)
    return templates.TemplateResponse("register.html", {"request": request, "departments": departments})


@app.post("/register", response_class=HTMLResponse)
async def register_post(
    request: Request,
    username: str = Form(...),
    password: str = Form(...),
    email: str = Form(...),
    role: str = Form(...),
    roll_number: Optional[str] = Form(None),
    register_number: Optional[str] = Form(None),
    department_id_str: Optional[str] = Form(None), 
    db: AsyncSession = Depends(get_db)
):

    password_hash = bcrypt.hashpw(
        password.encode('utf-8'),
        bcrypt.gensalt(rounds=12) 
    ).decode('utf-8')


    department_id: Optional[int] = None
    if role != "student":
        roll_number = register_number = None
        department_id_str = None

    if department_id_str:
        try:
            department_id = int(department_id_str)
        except (ValueError, TypeError):
            department_id = None

    if department_id is not None:
        exists = await db.scalar(
            select(Department.id).where(Department.id == department_id)
        )
        if not exists:
            departments = await db.scalars(
                select(Department.id, Department.name).order_by(Department.name)
            )
            dept_list = [{"id": d.id, "name": d.name} for d in departments.all()]
            return templates.TemplateResponse("register.html", {
                "request": request,
                "error": f"Department with id {department_id} does not exist.",
                "departments": dept_list,
                "form_data": request._form._dict 
            })

    email_exists = await db.scalar(
        select(User.id).where(func.lower(User.email) == email.strip().lower())
    )
    if email_exists:
        departments = await _get_departments_list(db)
        return templates.TemplateResponse("register.html", {
            "request": request,
            "error": "Email already exists.",
            "departments": departments,
            "form_data": request._form._dict
        })

    if role == "student":
        if roll_number:
            roll_exists = await db.scalar(
                select(User.id).where(User.roll_number == roll_number.strip())
            )
            if roll_exists:
                departments = await _get_departments_list(db)
                return templates.TemplateResponse("register.html", {
                    "request": request,
                    "error": "Roll number already exists.",
                    "departments": departments,
                    "form_data": request._form._dict
                })

        if register_number:
            reg_exists = await db.scalar(
                select(User.id).where(User.register_number == register_number.strip())
            )
            if reg_exists:
                departments = await _get_departments_list(db)
                return templates.TemplateResponse("register.html", {
                    "request": request,
                    "error": "Registration number already exists.",
                    "departments": departments,
                    "form_data": request._form._dict
                })

    new_user = User(
        username=username.strip(),
        password_hash=password_hash,
        email=email.strip().lower(),
        role=role,
        approved=0,  # pending approval
        roll_number=roll_number.strip() if roll_number else None,
        register_number=register_number.strip() if register_number else None,
        department_id=department_id
    )

    db.add(new_user)
    try:
        await db.commit()
    except Exception as e:
        await db.rollback()
        departments = await _get_departments_list(db)
        return templates.TemplateResponse("register.html", {
            "request": request,
            "error": "Registration failed. Please try again.",
            "departments": departments,
            "form_data": request._form._dict
        })

    return templates.TemplateResponse("login.html", {
        "request": request,
        "message": "Registration successful. Await admin approval."
    })


async def _get_departments_list(db: AsyncSession):
    result = await db.execute(
        select(Department.id, Department.name).order_by(Department.name)
    )
    return [{"id": d.id, "name": d.name} for d in result.all()]







def require_login(request: Request):
    if not request.session.get("user_id"):
        raise HTTPException(status_code=401, detail="Please log in to access this page.")





async def is_team_lead(user_id: int, db: AsyncSession) -> bool:
    result = await db.execute(
        select(User.is_team_lead)
        .where(User.id == user_id)
    )
    value = result.scalar_one_or_none()
    return bool(value == 1) 




async def process_exam_submission(
    attempt_id: int,
    questions: list[Question],
    form_data: dict,
    db: AsyncSession
):

    try:
        # Re-fetch attempt to avoid stale data
        result = await db.execute(
            select(ExamAttempt).where(ExamAttempt.id == attempt_id)
        )
        attempt = result.scalars().first()
        if not attempt or attempt.status == "submitted":
            return

        total_score = 0
        correct_count = 0
        responses_to_save = []

        question_map = {q.id: q for q in questions}

        for key, selected_option_id in form_data.items():
            if not key.startswith("question_"):
                continue
            try:
                question_id = int(key.split("_")[1])
                selected_option_id = int(selected_option_id)
            except (ValueError, IndexError):
                continue

            question = question_map.get(question_id)
            if not question:
                continue

            # Find correct option
            correct_option = next((opt for opt in question.options if opt.is_correct), None)
            is_correct = correct_option and correct_option.id == selected_option_id

            if is_correct:
                correct_count += 1
                total_score += 1  # or weighted score later

            responses_to_save.append(
                ExamResponse(
                    exam_attempt_id=attempt_id,
                    question_id=question_id,
                    selected_option_id=selected_option_id
                )
            )

        # Update attempt
        attempt.score = total_score
        attempt.status = "submitted"
        attempt.end_time = datetime.now(ZoneInfo("Asia/Kolkata"))

        # Bulk insert responses
        db.add_all(responses_to_save)

        await db.commit()

    except Exception as e:
        await db.rollback()
        # Log error in production
        print(f"Error processing submission for attempt {attempt_id}: {e}")

    
class FeedbackReplyRequest(BaseModel):
    feedback_id: int
    admin_id: int
    reply_message: str

class FeedbackRequest(BaseModel):
    user_id: int
    rating: int
    message: str



# === 1. Submit Feedback ===
@app.post("/api/feedback/")
async def submit_feedback(
    feedback: FeedbackRequest,
    db: AsyncSession = Depends(get_db)
):
    if not (1 <= feedback.rating <= 5):
        raise HTTPException(
            status_code=400,
            detail="Rating must be between 1 and 5"
        )

    # Optional: Verify user exists (security best practice)
    user_exists = await db.get(User, feedback.user_id)
    if not user_exists:
        raise HTTPException(status_code=404, detail="User not found")

    new_feedback = Feedback(
        user_id=feedback.user_id,
        rating=feedback.rating,
        message=feedback.message.strip()
    )

    db.add(new_feedback)
    await db.commit()
    await db.refresh(new_feedback)

    return {
        "message": "Feedback submitted successfully",
        "feedback_id": new_feedback.id
    }


# === 2. Get feedback by user (with replies) ===
@app.get("/api/feedback/user/{user_id}")
async def get_user_feedback(user_id: int, db: AsyncSession = Depends(get_db)):
    result = await db.execute(
        select(Feedback)
        .options(
            selectinload(Feedback.replies).selectinload(FeedbackReply.admin),
        )
        .where(Feedback.user_id == user_id)
        .order_by(Feedback.created_at.desc())
    )

    feedbacks = result.scalars().all()

    feedback_list = []
    for fb in feedbacks:
        feedback_dict = {
            "id": fb.id,
            "user_id": fb.user_id,
            "rating": fb.rating,
            "message": fb.message,
            "created_at": fb.created_at.isoformat() if fb.created_at else None,
            "feedback_replies": [
                {
                    "reply_message": reply.reply_message,
                    "created_at": reply.created_at.isoformat() if reply.created_at else None,
                }
                for reply in fb.replies
            ]
        }
        feedback_list.append(feedback_dict)

    return {"feedback": feedback_list}


@app.get("/api/feedback/admin/")
async def get_all_feedback(db: AsyncSession = Depends(get_db)):
    result = await db.execute(
        select(Feedback)
        .options(
            selectinload(Feedback.user),
            selectinload(Feedback.replies).selectinload(FeedbackReply.admin)
        )
        .order_by(Feedback.created_at.desc())
    )

    feedbacks = result.scalars().all()

    feedback_list = []
    for fb in feedbacks:
        feedback_dict = {
            "id": fb.id,
            "user_id": fb.user_id,
            "rating": fb.rating,
            "message": fb.message,
            "created_at": fb.created_at.isoformat() if fb.created_at else None,
            "users": {
                "id": fb.user.id,
                "username": fb.user.username
            },
            "feedback_replies": [
                {
                    "reply_message": r.reply_message,
                    "created_at": r.created_at.isoformat() if r.created_at else None,
                }
                for r in fb.replies
            ]
        }
        feedback_list.append(feedback_dict)

    return {"feedback": feedback_list}


@app.post("/api/feedback/reply/")
async def reply_to_feedback(
    reply: FeedbackReplyRequest,
    db: AsyncSession = Depends(get_db)
):
    # Check if feedback exists
    feedback_exists = await db.get(Feedback, reply.feedback_id)
    if not feedback_exists:
        raise HTTPException(status_code=404, detail="Feedback not found")

    # Optional: verify admin exists and has role == "admin"
    admin = await db.get(User, reply.admin_id)
    if not admin:
        raise HTTPException(status_code=404, detail="Admin user not found")
    if admin.role != "admin":
        raise HTTPException(status_code=403, detail="Only admins can reply to feedback")

    new_reply = FeedbackReply(
        feedback_id=reply.feedback_id,
        admin_id=reply.admin_id,
        reply_message=reply.reply_message.strip()
    )

    db.add(new_reply)
    await db.commit()
    await db.refresh(new_reply)

    return {"message": "Reply added successfully"}


@app.get("/feedback/", response_class=HTMLResponse)
async def feedback_page(request: Request, db: AsyncSession = Depends(get_db)):
    user_id_str = request.session.get("user_id")
    if not user_id_str:
        raise HTTPException(status_code=401, detail="Not authenticated")

    try:
        user_id = int(user_id_str)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid user ID in session")

    # Verify user exists
    user = await db.get(User, user_id)
    if not user:
        raise HTTPException(status_code=401, detail="User not found")

    return templates.TemplateResponse(
        "feedback.html",
        {"request": request, "user_id": user_id}
    )







@app.post("/api/notifications/send")
async def send_notification(
    request: Request,
    target: str = Form(...),
    message: str = Form(...),
    department_id: Optional[int] = Form(None),
    db: AsyncSession = Depends(get_db)
):
    user_id = request.session.get("user_id")
    user_role = request.session.get("role")

    if not user_id or user_role not in ["admin", "teacher"]:
        raise HTTPException(status_code=401, detail="Unauthorized")

    # === Determine valid targets based on role ===
    allowed_targets = ["all_students"]
    if user_role == "admin":
        allowed_targets.extend(["all_teachers", "department"])

    if target : target.strip()
    if target not in allowed_targets:
        raise HTTPException(status_code=400, detail="Invalid target")

    # === Build recipient query ===
    stmt = select(User.id).where(User.role == "student")

    if target == "all_teachers" and user_role == "admin":
        stmt = select(User.id).where(User.role == "teacher")
    elif target == "department":
        if not department_id:
            raise HTTPException(status_code=400, detail="department_id required for department target")
        stmt = select(User.id).where(
            User.role == "student",
            User.department_id == department_id
        )

    # Execute recipient query
    result = await db.execute(stmt)
    recipient_ids = [row[0] for row in result.all()]

    if not recipient_ids:
        raise HTTPException(status_code=404, detail="No recipients found for the selected target")

    # === Create Notification ===
    notif_stmt = insert(Notification).values(
        sender_id=int(user_id),
        message=message.strip()
    ).returning(Notification.id)

    notif_result = await db.execute(notif_stmt)
    notification_id = notif_result.scalar_one()

    # === Bulk insert recipients ===
    recipient_entries = [
        {"notification_id": notification_id, "recipient_id": rid}
        for rid in recipient_ids
    ]

    await db.execute(insert(NotificationRecipient), recipient_entries)
    await db.commit()

    return {"detail": "Notification sent successfully"}


@app.get("/notifications/new", response_class=HTMLResponse)
async def new_notification_page(
    request: Request,
    db: AsyncSession = Depends(get_db)
):
    user_id = request.session.get("user_id")
    user_role = request.session.get("role")

    if not user_id or user_role not in ["admin", "teacher"]:
        raise HTTPException(status_code=401, detail="Unauthorized")

    # === Fetch all departments (for dropdown) ===
    dept_result = await db.execute(
        select(Department.id, Department.name)
        .order_by(Department.name)
    )
    departments = [{"id": d.id, "name": d.name} for d in dept_result.all()]

    # === Fetch recent notifications with sender info ===
    notif_result = await db.execute(
        select(
            Notification.id,
            Notification.message,
            Notification.created_at,
            User.username.label("sender_name")
        )
        .join(User, User.id == Notification.sender_id)
        .order_by(Notification.created_at.desc())
        .limit(50)  # Prevent loading thousands
    )

    notifications = [
        {
            "id": n.id,
            "message": n.message,
            "created_at": n.created_at,
            "sender_name": n.sender_name,
        }
        for n in notif_result.all()
    ]

    return templates.TemplateResponse(
        "create_notification.html",
        {
            "request": request,
            "departments": departments,
            "user_role": user_role,
            "notifications": notifications,
            "is_admin": user_role == "admin",
        }
    )




@app.delete("/api/notifications/{notification_id}/delete")
async def delete_notification(
    notification_id: int,
    db: AsyncSession = Depends(get_db)
) -> Dict[str, str]:
    

    delete_recipients_stmt = delete(NotificationRecipient).where(
        NotificationRecipient.notification_id == notification_id
    )
    result_recipients = await db.execute(delete_recipients_stmt)

    # Delete the notification itself
    delete_notification_stmt = delete(Notification).where(
        Notification.id == notification_id
    )
    result_notification = await db.execute(delete_notification_stmt)

    if result_notification.rowcount == 0:
        raise HTTPException(
            status_code=404,
            detail="Notification not found or already deleted"
        )

    await db.commit()

    return {"detail": "Notification deleted successfully"}


@app.get("/api/notifications/count")
async def notifications_count(
    request: Request,
    db: AsyncSession = Depends(get_db)
) -> JSONResponse:

    user_id = request.session.get("user_id")
    if not user_id:
        raise HTTPException(status_code=401, detail="Unauthorized")

    try:
        user_id = int(user_id)
    except (ValueError, TypeError):
        raise HTTPException(status_code=400, detail="Invalid user ID in session")

    # Count unread notifications
    count_stmt = select(func.count()).select_from(NotificationRecipient)\
        .where(
            NotificationRecipient.recipient_id == user_id,
            NotificationRecipient.is_read.is_(False)
        )

    result = await db.execute(count_stmt)
    unread_count: int = result.scalar_one()

    return JSONResponse({"unread_count": unread_count})


@app.get("/api/notifications/view")
async def view_notifications(
    request: Request,
    db: AsyncSession = Depends(get_db)
) -> JSONResponse:

    user_id = request.session.get("user_id")
    if not user_id:
        raise HTTPException(status_code=401, detail="Unauthorized")

    try:
        user_id = int(user_id)
    except (ValueError, TypeError):
        raise HTTPException(status_code=400, detail="Invalid user ID in session")

    # 1. Mark all as read
    update_stmt = update(NotificationRecipient)\
        .where(NotificationRecipient.recipient_id == user_id)\
        .values(is_read=True)

    await db.execute(update_stmt)

    # 2. Fetch all notifications for user with sender username
    stmt = select(NotificationRecipient, Notification, User.username)\
        .join(Notification, Notification.id == NotificationRecipient.notification_id)\
        .join(User, User.id == Notification.sender_id)\
        .where(NotificationRecipient.recipient_id == user_id)\
        .order_by(Notification.created_at.desc())

    result = await db.execute(stmt)
    rows = result.all()

    notifications = []
    for recipient, notification, sender_username in rows:
        notifications.append({
            "id": recipient.id,
            "notification_id": recipient.notification_id,
            "recipient_id": recipient.recipient_id,
            "is_read": recipient.is_read,
            "notifications": {
                "id": notification.id,
                "sender_id": notification.sender_id,
                "message": notification.message,
                "created_at": notification.created_at.isoformat() if notification.created_at else None,
                "sender": {
                    "username": sender_username
                }
            }
        })

    return JSONResponse({"notifications": notifications})



@app.post("/generate_download_token")
async def generate_download_token(request: Request):
    user_id = request.session.get("user_id")
    if not user_id:
        raise HTTPException(status_code=401, detail="Unauthorized")
    token = f"token_{user_id}"
    request.session["download_token"] = token
    return {"download_token": token}



@app.post("/forgot-password")
async def send_otp(
    request: Request,
    email: str = Form(...),
    db: AsyncSession = Depends(get_db)
):
    # Normalize email
    email = email.strip().lower()

    # Check if user exists
    result = await db.execute(select(User.id).where(User.email == email))
    user_exists = result.scalar_one_or_none() is not None

    if not user_exists:
        request.session["error"] = "If this email is registered, you will receive an OTP."
        return RedirectResponse(url="/forgot-password", status_code=303)

    # Generate secure 6-digit OTP
    otp = str(random.SystemRandom().randint(100000, 999999))

    # Store in session (in session (short-lived, secure)
    request.session["otp"] = otp
    request.session["otp_email"] = email
    request.session["otp_created_at"] = datetime.now(ZoneInfo("UTC")).isoformat()

    # Send OTP via your existing email function
    subject = "Your Password Reset OTP"
    body = f"""
    <p>Your OTP for password reset is:</p>
    <h2 style="font-size: 24px; letter-spacing: 5px;">{otp}</h2>
    <p>It is valid for <strong>10 minutes</strong>.</p>
    <p>If you didn't request this, please ignore this email.</p>
    """
    send_email(to=email, subject=subject, body=body, is_html=True)

    request.session["success"] = "OTP sent to your email!"
    return RedirectResponse(url="/verify-otp", status_code=303)


# ─────────────────────────────────────────────────────────────────────────────
# 2. Reset Password (after OTP verification)
# ─────────────────────────────────────────────────────────────────────────────
@app.post("/reset-password")
async def reset_password(
    request: Request,
    password: str = Form(...),
    confirm_password: str = Form(...),
    db: AsyncSession = Depends(get_db)
):
    if password != confirm_password:
        request.session["error"] = "Passwords do not match."
        return RedirectResponse(url="/reset-password", status_code=303)

    if len(password) < 8:
        request.session["error"] = "Password must be at least 8 characters."
        return RedirectResponse(url="/reset-password", status_code=303)

    email = request.session.get("otp_email")
    if not email:
        request.session["error"] = "Session expired. Please request a new OTP."
        return RedirectResponse(url="/forgot-password", status_code=303)

    # Hash password securely with bcrypt
    password_hash = bcrypt.hashpw(
        password.encode("utf-8"),
        bcrypt.gensalt(rounds=12)  # 12 is standard for production
    ).decode("utf-8")

    # Update password in DB
    result = await db.execute(
        update(User)
        .where(User.email == email)
        .values(password_hash=password_hash)
    )

    if result.rowcount == 0:
        request.session["error"] = "Failed to update password. Try again."
        return RedirectResponse(url="/reset-password", status_code=303)

    await db.commit()

    # Clear session data
    for key in ["otp", "otp_email", "otp_created_at"]:
        request.session.pop(key, None)

    # Send confirmation
    send_email(
        to=email,
        subject="Password Changed Successfully",
        body="""
        <p>Your password has been successfully updated.</p>
        <p>If this wasn't you, please contact support immediately.</p>
        """,
        is_html=True
    )

    request.session["success"] = "Password reset successful! You can now log in."
    return RedirectResponse(url="/login", status_code=303)


@app.get("/exam_result", response_class=HTMLResponse)
async def exam_result_form(
    request: Request,
    db: AsyncSession = Depends(get_db)
):
    user_id = request.session.get("user_id")
    if not user_id:
        raise HTTPException(status_code=401, detail="Authentication required")

    try:
        user_id = int(user_id)
    except (ValueError, TypeError):
        raise HTTPException(status_code=400, detail="Invalid session")

    # Only teachers should access this
    result = await db.execute(
        select(User.role).where(User.id == user_id)
    )
    role = result.scalar_one_or_none()

    if role != "teacher":
        raise HTTPException(status_code=403, detail="Only teachers can view exam results")

    # Fetch all exams created by this teacher
    exams_result = await db.execute(
        select(Exam.id, Exam.name, Exam.subject, Exam.start_time)
        .where(Exam.teacher_id == user_id)
        .order_by(Exam.start_time.desc())
    )

    exams = [
        {
            "id": exam_id,
            "name": name,
            "subject": subject,
            "start_time": start_time,
        }
        for exam_id, name, subject, start_time in exams_result.all()
    ]

    return templates.TemplateResponse(
        "exam_result_form.html",
        {
            "request": request,
            "exams": exams,
        }
    )


header_mapping = {
    "roll_number": "Roll No",
    "register_number": "Reg No",
    "username": "Name",
    "score": "Score",
    "percentage": "Percentage"
}



def validate_download_token(request: Request, token: str):
    expected = request.session.get("download_token")
    if not expected or expected != token:
        raise HTTPException(status_code=403, detail="Invalid download token")
    # Clear after use
    request.session.pop("download_token", None)



@app.api_route("/exam_result/download", methods=["GET", "POST"])
async def download_exam_result(
    request: Request,
    db: AsyncSession = Depends(get_db),
    exam_id: int = Query(...),
    exam_name: Optional[str] = Query(None),
    orderby: Optional[str] = Query(None),
    file_type: Optional[str] = Query("pdf", regex="^(pdf|docx)$"),
    include_roll_no: bool = Query(False),
    include_reg_no: bool = Query(False),
    include_username: bool = Query(False),
    include_score: bool = Query(True),
    include_percentage: bool = Query(True),
    download_token: Optional[str] = Query(None),
):

    require_login(request)
    user_id = request.session.get("user_id")
    role = request.session.get("role")

    if role not in ["teacher", "admin"]:
        raise HTTPException(status_code=403, detail="Not authorized")

    # Optional: validate one-time download token
    if download_token:
        validate_download_token(request, download_token)

    # === Fetch Exam + Validate Ownership ===
    exam_result = await db.execute(
        select(Exam)
        .where(Exam.id == exam_id)
    )
    exam: Exam = exam_result.scalars().first()

    if not exam:
        raise HTTPException(status_code=404, detail="Exam not found")

    # Teachers can only download their own exams
    if role == "teacher" and exam.teacher_id != user_id:
        raise HTTPException(status_code=403, detail="You can only download results for your own exams")

    # === Get total questions (total marks) ===
    total_questions_result = await db.execute(
        select(func.count(Question.id))
        .where(Question.exam_id == exam_id)
    )
    total_marks = total_questions_result.scalar_one()

    if total_marks == 0:
        raise HTTPException(status_code=400, detail="Exam has no questions")

    # === Get all students in the department ===
    students_result = await db.execute(
        select(User.id, User.roll_number, User.register_number, User.username)
        .where(
            and_(
                User.role == "student",
                User.department_id == exam.department_id
            )
        )
        .order_by(User.username)
    )
    students = students_result.all()  # List[Row]

    # === Get all attempts for this exam (student_id → score) ===
    attempts_result = await db.execute(
        select(ExamAttempt.student_id, ExamAttempt.score)
        .where(ExamAttempt.exam_id == exam_id)
    )
    attempts_dict = {row.student_id: row.score for row in attempts_result.all()}

    # === Build result rows dynamically ===
    result_list: List[Dict[str, Any]] = []

    for student_id, roll_no, reg_no, username in students:
        score = attempts_dict.get(student_id)
        obtained = score if score is not None else 0
        percentage = (obtained / total_marks * 100) if total_marks > 0 else 0

        row: Dict[str, Any] = {}
        if include_roll_no:
            row["roll_number"] = roll_no or ""
        if include_reg_no:
            row["register_number"] = reg_no or ""
        if include_username:
            row["username"] = username or "Unknown"
        if include_score:
            row["score"] = f"{obtained}/{total_marks}"
        if include_percentage:
            row["percentage"] = f"{percentage:.2f}%"

        if row:  # Only add if at least one field is included
            result_list.append(row)

    # === Sorting ===
    valid_sort_keys = {"roll_number", "register_number", "score", "percentage", "username"}
    if orderby in valid_sort_keys:
        # Extract numeric score for proper sorting when sorting by score/percentage
        if orderby in ["score", "percentage"]:
            def sort_key(item):
                val = item.get(orderby, "0")
                try:
                    return float(val.split("/")[0] if "/" in val else val.rstrip("%"))
                except:
                    return 0
            result_list.sort(key=sort_key, reverse=True)
        else:
            result_list.sort(key=lambda x: str(x.get(orderby, "")).lower())

    exam_title = exam_name or exam.name or "Exam Results"
    safe_exam_name = "".join(c if c.isalnum() or c in " _-" else "_" for c in exam_title)

    # === Generate PDF ===
    if file_type.lower() == "pdf":
        try:
            class PDF(FPDF):
                def header(self):
                    self.set_font("Arial", "B", 16)
                    self.cell(0, 10, exam_title, ln=True, align="C")
                    self.ln(10)

                def footer(self):
                    self.set_y(-15)
                    self.set_font("Arial", "I", 8)
                    self.cell(0, 10, "Made with Testify", align="L")
                    self.cell(0, 10, f"Page {self.page_no()}", align="R")

            pdf = PDF()
            pdf.add_page()
            pdf.set_font("Arial", "", 12)

            if not result_list:
                pdf.cell(0, 10, "No results to display.", ln=True)
            else:
                headers = list(result_list[0].keys())
                col_width = 190 / len(headers)  # A4 width ~210mm, margin 10 each side
                line_height = 8

                # Header
                pdf.set_font("Arial", "B", 12)
                for h in headers:
                    pdf.cell(col_width, line_height, header_mapping.get(h, h.replace("_", " ").title()), border=1, align="C")
                pdf.ln(line_height)

                # Rows
                pdf.set_font("Arial", "", 11)
                for row in result_list:
                    # Handle text wrapping
                    cells = []
                    max_lines = 1
                    for h in headers:
                        text = str(row.get(h, ""))
                        wrapped = textwrap.wrap(text, width=int(col_width / 2.5))
                        cells.append(wrapped or [""])
                        max_lines = max(max_lines, len(wrapped))

                    cell_height = line_height * max_lines

                    if pdf.get_y() + cell_height > pdf.page_break_trigger:
                        pdf.add_page()
                        pdf.set_font("Arial", "B", 12)
                        for h in headers:
                            pdf.cell(col_width, line_height, header_mapping.get(h, h.replace("_", " ").title()), border=1, align="C")
                        pdf.ln(line_height)
                        pdf.set_font("Arial", "", 11)

                    for i in range(max_lines):
                        for cell_lines in cells:
                            text = cell_lines[i] if i < len(cell_lines) else ""
                            pdf.cell(col_width, line_height, text, border=1, align="C")
                        pdf.ln(line_height)
                    pdf.ln(2)

            temp_file = NamedTemporaryFile(delete=False, suffix=".pdf")
            pdf.output(temp_file.name)
            temp_file.close()

            return FileResponse(
                temp_file.name,
                media_type="application/pdf",
                filename=f"{safe_exam_name}_Results.pdf",
                background=None  # Important: prevents early deletion
            )

        except ImportError:
            raise HTTPException(status_code=500, detail="FPDF not installed. Run: pip install fpdf2")

    # === Generate DOCX ===
    elif file_type.lower() == "docx":
        try:
            document = Document()
            document.add_heading(exam_title, 0)

            if not result_list:
                document.add_paragraph("No results to display.")
            else:
                headers = list(result_list[0].keys())
                table = document.add_table(rows=1, cols=len(headers))
                table.style = "Table Grid"

                hdr_cells = table.rows[0].cells
                for i, key in enumerate(headers):
                    hdr_cells[i].text = header_mapping.get(key, key.replace("_", " ").title())

                for row_data in result_list:
                    row_cells = table.add_row().cells
                    for i, key in enumerate(headers):
                        row_cells[i].text = str(row_data.get(key, ""))

            # Footer
            section = document.sections[0]
            footer = section.footer
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.text = "Made with Testify"

            temp_file = NamedTemporaryFile(delete=False, suffix=".docx")
            document.save(temp_file.name)
            temp_file.close()

            return FileResponse(
                temp_file.name,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                filename=f"{safe_exam_name}_Results.docx"
            )

        except ImportError:
            raise HTTPException(status_code=500, detail="python-docx not installed. Run: pip install python-docx")

    else:
        raise HTTPException(status_code=400, detail="Unsupported file type. Use 'pdf' or 'docx'")


def send_email(to_email: str, subject: str, body: str):
    try:
        msg = MIMEMultipart("alternative")
        msg["From"] = EMAIL_USER
        msg["To"] = to_email
        msg["Subject"] = subject
        
        text = """Welcome"""
        html = body
        
        # Attach both versions
        msg.attach(MIMEText(text, "plain"))
        msg.attach(MIMEText(html, "html"))
        
        with smtplib.SMTP("smtp.gmail.com", 587) as server:
            server.starttls()
            server.login(EMAIL_USER, EMAIL_PASS)
            server.sendmail(EMAIL_USER, to_email, msg.as_string())
            
    except smtplib.SMTPException as e:
        print(f"SMTP error sending to {to_email}: {e}")
    except Exception as e:
        print(f"General error sending to {to_email}: {e}")


@app.get("/forgot-password")
async def forgot_password_form(request: Request):
    return templates.TemplateResponse("forgot_password.html", {"request": request})



@app.get("/verify-otp")
async def verify_otp_form(request: Request):
    return templates.TemplateResponse("verify_otp.html", {"request": request})

@app.post("/verify-otp")
async def verify_otp(request: Request, otp: str = Form(...)):
    stored_otp = request.session.get("otp")
    if otp != stored_otp:
        request.session["error"] = "Invalid OTP"
        return RedirectResponse(url="/verify-otp", status_code=303)
    return RedirectResponse(url="/reset-password", status_code=303)

@app.get("/reset-password")
async def reset_password_form(request: Request):
    return templates.TemplateResponse("reset_password.html", {"request": request})



@app.get("/download-release")
async def forgot_password_form(request: Request):
    return templates.TemplateResponse("release_page.html", {"request": request})
